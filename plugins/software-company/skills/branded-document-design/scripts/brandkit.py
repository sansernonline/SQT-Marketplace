"""
brandkit.py — ตัวช่วยสร้างเอกสาร Word (.docx) ที่มีหน้าตาเป็นระบบเดียวกัน
สไตล์อ้างอิงจาก Apps Track SRS (ฟอนต์ Tahoma · accent น้ำเงิน-ม่วง · ตารางหัวฟ้าอ่อน)

ใช้กับ python-docx >= 1.1
    pip install python-docx

ตัวอย่างสั้นที่สุด
------------------
    from brandkit import BrandDoc, TOKENS

    doc = BrandDoc()                       # A4 · Tahoma · โทน Apps Track
    doc.cover("เอกสารข้อกำหนดซอฟต์แวร์ (Software Specification)",
              subtitle="ระบบ Apps Track — Project Control & Monitor",
              meta="เวอร์ชันเอกสาร 3.5  •  ปรับปรุง 19 กรกฎาคม 2026",
              logo="asset/AppsTrack_Logo_Badge.png")
    doc.h1("1. ภาพรวมระบบ")
    doc.para("eitprojects เป็นระบบบริหารและติดตามโครงการ ...")
    doc.table(["หัวข้อ", "รายละเอียด"],
              [["URL ระบบ", "https://project.eitaccount.cloud"],
               ["ประเภทระบบ", "Web Application (SPA) + PWA"]],
              widths=[2600, 6760])
    doc.callout("tip", "ข้อแนะนำ", "เขียน callout สั้น ๆ 1–3 บรรทัดเท่านั้น")
    doc.save("out.docx")

ทุกเมธอดคืนค่าอ็อบเจกต์ที่สร้าง จึงต่อยอด/ปรับแต่งเพิ่มเองได้เสมอ
"""

from __future__ import annotations

import copy
from dataclasses import dataclass, field

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ---------------------------------------------------------------------------
# 1 · Design tokens — แก้ที่เดียว เปลี่ยนทั้งเอกสาร
# ---------------------------------------------------------------------------

TOKENS = {
    # แบรนด์
    "brand": "2A78D6",        # น้ำเงินหลัก — หัวข้อ H1, ตัวเลข KPI, ลิงก์
    "brand_2": "6A5CD6",      # ม่วง — accent รอง, แถบ callout
    "brand_deep": "2A4C86",   # น้ำเงินเข้ม — H2, หัวตาราง
    "brand_tint": "EDF1FB",   # ฟ้าอ่อนมาก — พื้นหัวตาราง, พื้น KPI
    "brand_tint_2": "F6F8FD", # ฟ้าอ่อนกว่า — แถวสลับในตารางยาว
    # ตัวอักษร
    "text": "333B4A",         # หัวข้อย่อย/ข้อความเน้น
    "text_body": "414957",    # เนื้อความ
    "text_muted": "7D8492",   # คำบรรยายรูป, meta, footer
    "text_faint": "A9AEB9",
    # เส้น
    "line": "E4E7EE",
    # สถานะ (ใช้กับ pill / callout / ตารางสรุป)
    "green": "17794A", "green_bg": "E9F7EF",
    "blue": "2160AB",  "blue_bg": "EAF2FD",
    "amber": "96660D", "amber_bg": "FDF5E4",
    "red": "A63A34",   "red_bg": "FDEDEC",
    "violet": "52439F", "violet_bg": "F1EEFC",
    "grey": "626A7A",  "grey_bg": "F2F4F8",
}

# ฟอนต์: Tahoma ปลอดภัยที่สุดสำหรับเอกสารไทย+อังกฤษบน Windows/Word
# (มีทุกเครื่อง · วรรณยุกต์ไม่ลอย · น้ำหนัก bold อ่านง่าย)
FONT = "Tahoma"

# ขนาด (pt)
SIZE = {
    "cover_title": 20, "cover_subtitle": 14.5, "cover_meta": 9, "cover_note": 8,
    "h1": 16, "h2": 12.5, "h3": 11.5,
    "body": 11, "table": 11, "small": 9, "caption": 8.5, "footer": 9,
}

CALLOUT_KINDS = {
    #  key   : (สี bg, สีขอบซ้าย, สีตัวอักษรหัวข้อ, สัญลักษณ์)
    "tip":     ("blue_bg",   "brand",    "blue",   "💡"),
    "note":    ("grey_bg",   "text_muted", "grey", "ℹ️"),
    "warning": ("amber_bg",  "amber",    "amber",  "⚠️"),
    "critical":("red_bg",    "red",      "red",    "🚨"),
    "success": ("green_bg",  "green",    "green",  "✅"),
    "question":("violet_bg", "brand_2",  "violet", "❓"),
}


# ---------------------------------------------------------------------------
# 2 · ตัวช่วยระดับ XML
# ---------------------------------------------------------------------------

def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn("w:" + k), str(v))
    return e


def style_run(run, *, size=None, color=None, bold=False, italic=False,
              font=FONT, mono=False):
    """ตั้งฟอนต์ให้ครบทั้ง latin / complex-script — จำเป็นสำหรับภาษาไทย
    ถ้าตั้งแต่ run.font.size อย่างเดียว ตัวอักษรไทยจะไม่เปลี่ยนขนาดตาม
    เพราะ Word ถือว่าไทยเป็น complex script (ต้องตั้ง szCs / bCs / iCs ด้วย)"""
    name = "Consolas" if mono else font
    run.font.name = name
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = _el("w:rFonts")
        rpr.insert(0, rfonts)
    for a in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn("w:" + a), name)
    if size is not None:
        half = str(int(round(size * 2)))
        for tag in ("w:sz", "w:szCs"):
            rpr.append(_el(tag, val=half))
    if bold:
        rpr.append(_el("w:b"))
        rpr.append(_el("w:bCs"))
    if italic:
        rpr.append(_el("w:i"))
        rpr.append(_el("w:iCs"))
    if color:
        run.font.color.rgb = RGBColor.from_string(TOKENS.get(color, color))
    order_rpr(rpr)
    return run


# ลำดับ element ภายใน <w:rPr> ตามสเปก ECMA-376 — Word จะเมินค่าที่วางผิดลำดับ
_RPR_ORDER = ["rStyle", "rFonts", "b", "bCs", "i", "iCs", "caps", "smallCaps",
              "strike", "dstrike", "outline", "shadow", "emboss", "imprint",
              "noProof", "snapToGrid", "vanish", "webHidden", "color", "spacing",
              "w", "kern", "position", "sz", "szCs", "highlight", "u", "effect",
              "bdr", "shd", "fitText", "vertAlign", "rtl", "cs", "em", "lang",
              "eastAsianLayout", "specVanish", "oMath"]


def order_rpr(rpr):
    """เรียง child ของ rPr ใหม่ให้ถูกลำดับ — จำเป็นเมื่อเราแทรก element เอง"""
    idx = {qn("w:" + t): i for i, t in enumerate(_RPR_ORDER)}
    children = list(rpr)
    for c in children:
        rpr.remove(c)
    for c in sorted(children, key=lambda e: idx.get(e.tag, 999)):
        rpr.append(c)


def shade(cell, hex_or_token):
    """ระบายพื้นหลังเซลล์ตาราง"""
    fill = TOKENS.get(hex_or_token, hex_or_token)
    cell._tc.get_or_add_tcPr().append(_el("w:shd", val="clear", color="auto", fill=fill))


def cell_margins(table, top=50, left=90, bottom=50, right=90):
    """ระยะขอบในเซลล์ (หน่วย twips · 20 twips = 1pt) — ค่าเริ่มต้นของ Word แน่นเกินไป"""
    tblpr = table._tbl.tblPr
    mar = _el("w:tblCellMar")
    for tag, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        mar.append(_el("w:" + tag, w=val, type="dxa"))
    tblpr.append(mar)


def set_borders(table, color="line", size=4, inside=True):
    """เส้นตาราง: บาง สีเทาอ่อน — ไม่ใช่เส้นดำหนาแบบ default"""
    col = TOKENS.get(color, color)
    tblpr = table._tbl.tblPr
    borders = _el("w:tblBorders")
    edges = ["top", "left", "bottom", "right"] + (["insideH", "insideV"] if inside else [])
    for edge in edges:
        borders.append(_el("w:" + edge, val="single", sz=size, space=0, color=col))
    tblpr.append(borders)


def fixed_widths(table, widths):
    """บังคับความกว้างคอลัมน์ให้ตรงตามที่กำหนด (หน่วย twips)
    python-docx อย่างเดียวไม่พอ — Word จะคำนวณใหม่ถ้าไม่ปิด autofit + ไม่ตั้ง tblGrid"""
    tbl = table._tbl
    tblpr = tbl.tblPr
    tblpr.append(_el("w:tblLayout", type="fixed"))
    tblpr.append(_el("w:tblW", w=str(sum(widths)), type="dxa"))
    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        tbl.remove(grid)
    grid = _el("w:tblGrid")
    for w in widths:
        grid.append(_el("w:gridCol", w=str(w)))
    tbl.insert(list(tbl).index(tblpr) + 1, grid)
    table.autofit = False
    for row in table.rows:
        for ci, w in enumerate(widths):
            if ci < len(row.cells):
                row.cells[ci].width = Pt(w / 20)


def no_borders(table):
    tblpr = table._tbl.tblPr
    borders = _el("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borders.append(_el("w:" + edge, val="none", sz=0, space=0, color="auto"))
    tblpr.append(borders)


def left_accent(cell, color="brand", size=24):
    """แถบสีหนาด้านซ้ายของเซลล์ — ใช้ทำ callout"""
    tcpr = cell._tc.get_or_add_tcPr()
    borders = _el("w:tcBorders")
    borders.append(_el("w:left", val="single", sz=size, space=0,
                       color=TOKENS.get(color, color)))
    tcpr.append(borders)


def repeat_header(row):
    """ให้แถวหัวตารางซ้ำทุกหน้าเวลาตารางยาวข้ามหน้า"""
    trpr = row._tr.get_or_add_trPr()
    trpr.append(_el("w:tblHeader", val="true"))


def keep_with_next(paragraph):
    """กันหัวข้อค้างท้ายหน้าโดยที่เนื้อหาไปอยู่หน้าถัดไป"""
    paragraph.paragraph_format.keep_with_next = True


def add_field(paragraph, instr, **run_kwargs):
    """แทรก field code ของ Word (PAGE, NUMPAGES, TOC ...)"""
    run = paragraph.add_run()
    style_run(run, **run_kwargs)
    r = run._r
    fld_begin = _el("w:fldChar", fldCharType="begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    fld_end = _el("w:fldChar", fldCharType="end")
    r.append(fld_begin)
    r.append(instr_el)
    r.append(fld_end)
    return run


# ---------------------------------------------------------------------------
# 3 · ตัวสร้างเอกสาร
# ---------------------------------------------------------------------------

@dataclass
class BrandDoc:
    """เอกสาร Word ที่ตั้งค่าสไตล์ให้เรียบร้อยแล้ว"""

    path_template: str | None = None
    page: str = "A4"                  # "A4" | "Letter"
    margins_cm: tuple = (2.2, 2.0, 2.2, 2.0)   # บน ขวา ล่าง ซ้าย
    footer_text: str = "หน้า"          # ข้อความนำหน้าเลขหน้า ("" = ไม่ใส่คำนำหน้า)
    doc: Document = field(init=False)

    def __post_init__(self):
        self.doc = Document(self.path_template) if self.path_template else Document()
        self._setup_page()
        self._setup_normal_style()
        self._setup_footer()

    # -- ตั้งค่าเริ่มต้น ---------------------------------------------------
    def _setup_page(self):
        s = self.doc.sections[0]
        if self.page == "A4":
            s.page_width, s.page_height = Cm(21.0), Cm(29.7)
        else:
            s.page_width, s.page_height = Cm(21.59), Cm(27.94)
        top, right, bottom, left = self.margins_cm
        s.top_margin, s.right_margin = Cm(top), Cm(right)
        s.bottom_margin, s.left_margin = Cm(bottom), Cm(left)

    def _setup_normal_style(self):
        st = self.doc.styles["Normal"]
        st.font.name = FONT
        st.font.size = Pt(SIZE["body"])
        st.font.color.rgb = RGBColor.from_string(TOKENS["text_body"])
        rpr = st.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        for a in ("ascii", "hAnsi", "cs", "eastAsia"):
            rfonts.set(qn("w:" + a), FONT)
        rpr.append(_el("w:szCs", val=str(int(SIZE["body"] * 2))))
        pf = st.paragraph_format
        pf.space_after = Pt(6)
        pf.line_spacing = 1.32          # เอกสารไทยต้องการระยะบรรทัดมากกว่าอังกฤษ

    def _setup_footer(self):
        p = self.doc.sections[0].footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if self.footer_text:
            style_run(p.add_run(self.footer_text + " "),
                      size=SIZE["footer"], color="text_muted")
        add_field(p, "PAGE", size=SIZE["footer"], color="text_muted")

    def _spacer(self, pt=5):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        style_run(p.add_run(""), size=pt)
        return p

    # -- บล็อกระดับหน้า ----------------------------------------------------
    def cover(self, title, subtitle=None, meta=None, note=None, logo=None,
              logo_width_cm=2.6, top_space_pt=150, page_break=True):
        """หน้าปก: โลโก้กลาง → ชื่อเอกสารสีแบรนด์ → ชื่อระบบ → บรรทัด meta → หมายเหตุ"""
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(top_space_pt)

        if logo:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(logo, width=Cm(logo_width_cm))
            p.paragraph_format.space_after = Pt(14)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        style_run(p.add_run(title), size=SIZE["cover_title"], color="brand", bold=True)

        if subtitle:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(18)
            style_run(p.add_run(subtitle), size=SIZE["cover_subtitle"],
                      color="text", bold=True)
        if meta:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            style_run(p.add_run(meta), size=SIZE["cover_meta"], color="text_body")
        if note:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_run(p.add_run(note), size=SIZE["cover_note"],
                      color="text_muted", italic=True)
        if page_break:
            self.doc.add_page_break()
        return self

    def toc(self, heading="สารบัญ", levels="1-3"):
        """แทรก field สารบัญ — ผู้ใช้กด F9 ใน Word เพื่ออัปเดต
        (LibreOffice/Word จะสร้างรายการให้อัตโนมัติเมื่อ refresh)"""
        self.h1(heading, numbered=False)
        p = self.doc.add_paragraph()
        add_field(p, f'TOC \\o "{levels}" \\h \\z \\u', size=SIZE["body"])
        self.doc.add_page_break()
        return p

    # -- หัวข้อและเนื้อความ -------------------------------------------------
    def h1(self, text, numbered=True, space_before=18):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(8)
        keep_with_next(p)
        style_run(p.add_run(text), size=SIZE["h1"], color="brand", bold=True)
        self._tag_outline(p, 1)
        return p

    def h2(self, text, space_before=14):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(6)
        keep_with_next(p)
        style_run(p.add_run(text), size=SIZE["h2"], color="brand_deep", bold=True)
        self._tag_outline(p, 2)
        return p

    def h3(self, text, space_before=10):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(4)
        keep_with_next(p)
        style_run(p.add_run(text), size=SIZE["h3"], color="text", bold=True)
        self._tag_outline(p, 3)
        return p

    def _tag_outline(self, paragraph, level):
        """ใส่ outline level เพื่อให้ TOC / Navigation pane เห็นหัวข้อ
        โดยไม่ต้องใช้ built-in Heading style (ซึ่งจะทับสีที่เราตั้งไว้)"""
        ppr = paragraph._p.get_or_add_pPr()
        ppr.append(_el("w:outlineLvl", val=str(level - 1)))

    def para(self, text=None, size=None, color="text_body", bold=False,
             italic=False, align=None, space_after=6):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        if align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align == "justify":
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if text:
            style_run(p.add_run(text), size=size or SIZE["body"],
                      color=color, bold=bold, italic=italic)
        return p

    def rich(self, parts, align=None, space_after=6):
        """ย่อหน้าที่ผสมหลายรูปแบบในบรรทัดเดียว
        parts = [("ข้อความ", {"bold": True, "color": "brand"}), ("ต่อ", {})]"""
        p = self.para(None, align=align, space_after=space_after)
        for text, opts in parts:
            style_run(p.add_run(text), size=opts.pop("size", SIZE["body"]), **opts)
        return p

    def bullets(self, items, style="List Bullet", size=None):
        out = []
        for it in items:
            p = self.doc.add_paragraph(style=style)
            p.paragraph_format.space_after = Pt(2)
            style_run(p.add_run(it), size=size or SIZE["body"], color="text_body")
            out.append(p)
        return out

    def code(self, text, lang=None):
        """บล็อกโค้ด — พื้นเทาอ่อน ฟอนต์ monospace"""
        t = self.doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        no_borders(t)
        cell_margins(t, 100, 140, 100, 140)
        c = t.cell(0, 0)
        shade(c, "grey_bg")
        left_accent(c, "line", 12)
        c.paragraphs[0].text = ""
        first = True
        for line in text.rstrip("\n").split("\n"):
            p = c.paragraphs[0] if first else c.add_paragraph()
            first = False
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            style_run(p.add_run(line), size=SIZE["small"], color="text", mono=True)
        self._spacer(6)
        return t

    # -- ตาราง --------------------------------------------------------------
    def table(self, headers, rows, widths=None, zebra=False, align=None,
              first_col_bold=False, size=None):
        """ตารางมาตรฐาน: หัวพื้นฟ้าอ่อน ตัวหนังสือน้ำเงินเข้ม เส้นเทาบาง
        widths — รายการความกว้างหน่วย twips (รวมกันประมาณ 9360 สำหรับ A4 margin 2cm)
        align  — รายการ 'left'|'center'|'right' ต่อคอลัมน์"""
        size = size or SIZE["table"]
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.autofit = widths is None
        set_borders(t)
        cell_margins(t)

        hdr = t.rows[0]
        repeat_header(hdr)
        for i, text in enumerate(headers):
            c = hdr.cells[i]
            shade(c, "brand_tint")
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            style_run(p.add_run(str(text)), size=size, color="brand_deep", bold=True)

        for ri, row in enumerate(rows):
            cells = t.add_row().cells
            for ci, val in enumerate(row):
                c = cells[ci]
                if zebra and ri % 2 == 1:
                    shade(c, "brand_tint_2")
                p = c.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                if align and ci < len(align):
                    p.alignment = {"center": WD_ALIGN_PARAGRAPH.CENTER,
                                   "right": WD_ALIGN_PARAGRAPH.RIGHT,
                                   "left": WD_ALIGN_PARAGRAPH.LEFT}[align[ci]]
                style_run(p.add_run("" if val is None else str(val)), size=size,
                          color="text_body", bold=first_col_bold and ci == 0)

        if widths:
            fixed_widths(t, widths)
        self._spacer(5)
        return t

    def kpi_row(self, items, width_twips=9360):
        """แถบตัวเลขสรุปแนวนอน — items = [("19", "โครงการทั้งหมด"), ...]"""
        t = self.doc.add_table(rows=1, cols=len(items))
        no_borders(t)
        cell_margins(t, 120, 120, 120, 120)
        for i, (value, label) in enumerate(items):
            c = t.cell(0, i)
            shade(c, "brand_tint")
            c.width = Pt(width_twips / 20 / len(items))
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            style_run(p.add_run(str(value)), size=18, color="brand", bold=True)
            p2 = c.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_after = Pt(0)
            style_run(p2.add_run(label), size=SIZE["caption"], color="text_muted")
        fixed_widths(t, [width_twips // len(items)] * len(items))
        self._spacer(7)
        return t

    # -- องค์ประกอบเสริม ----------------------------------------------------
    def callout(self, kind, title, body):
        """กล่องข้อความ: tip | note | warning | critical | success | question"""
        bg, accent, fg, icon = CALLOUT_KINDS[kind]
        t = self.doc.add_table(rows=1, cols=1)
        no_borders(t)
        cell_margins(t, 110, 150, 110, 150)
        c = t.cell(0, 0)
        shade(c, bg)
        left_accent(c, accent, 24)
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        style_run(p.add_run(f"{icon} {title}  "), size=SIZE["body"], color=fg, bold=True)
        style_run(p.add_run(body), size=SIZE["body"], color="text_body")
        fixed_widths(t, [9360])
        self._spacer(6)
        return t

    def figure(self, image_path, caption=None, width_cm=15.5, number=None):
        """รูปกลางหน้า + คำบรรยายตัวเอียงสีเทาใต้รูป"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        p.add_run().add_picture(image_path, width=Cm(width_cm))
        if caption:
            cap = self.doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(12)
            label = f"รูปที่ {number} — " if number else ""
            style_run(cap.add_run(label + caption), size=SIZE["caption"],
                      color="text_muted", italic=True)
        return p

    def pill_table(self, headers, rows, status_col, palette=None, widths=None):
        """ตารางที่คอลัมน์สถานะถูกย้อมสีตามค่า
        palette = {"เสร็จ": "green", "กำลังทำ": "amber", "ค้าง": "red"}"""
        palette = palette or {}
        t = self.table(headers, rows, widths=widths)
        for r in t.rows[1:]:
            cell = r.cells[status_col]
            key = cell.text.strip()
            tone = palette.get(key)
            if not tone:
                continue
            shade(cell, TOKENS[tone + "_bg"])
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                style_run(run, size=SIZE["table"], color=tone, bold=True)
        return t

    def signoff(self, roles, title="การลงนามอนุมัติ"):
        """ตารางเซ็นอนุมัติท้ายเอกสาร — roles = [("Product Owner", "คุณ..."), ...]"""
        self.h2(title)
        return self.table(["บทบาท", "ชื่อ-นามสกุล", "ลายเซ็น", "วันที่"],
                          [[r, n, "", ""] for r, n in roles],
                          widths=[2400, 2900, 2400, 1660])

    def landscape_section(self):
        """เปิดส่วนแนวนอน สำหรับตารางกว้างหรือแผนภาพใหญ่"""
        s = self.doc.add_section(WD_SECTION.NEW_PAGE)
        s.page_width, s.page_height = s.page_height, s.page_width
        return s

    def page_break(self):
        self.doc.add_page_break()

    def save(self, path):
        self.doc.save(path)
        return path


# ---------------------------------------------------------------------------
# 4 · ยูทิลิตี
# ---------------------------------------------------------------------------

def to_pdf(docx_path, outdir="."):
    """แปลงเป็น PDF ด้วย LibreOffice (ต้องมี soffice ในเครื่อง)
    ใช้ตรวจงานด้วยสายตาก่อนส่งมอบเสมอ"""
    import subprocess
    subprocess.run(["soffice", "--headless", "--convert-to", "pdf",
                    "--outdir", outdir, docx_path], check=True,
                   stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    return docx_path.rsplit("/", 1)[-1].replace(".docx", ".pdf")


def use_brand(**overrides):
    """เปลี่ยน palette ทั้งชุดสำหรับลูกค้า/โปรเจกต์อื่น
        use_brand(brand="C1121F", brand_deep="780000", brand_tint="FDECEC")"""
    TOKENS.update({k: v.lstrip("#").upper() for k, v in overrides.items()})
    return TOKENS
